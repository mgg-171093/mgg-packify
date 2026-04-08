"""
Tests para doc_generator.py

REQ-CAT-02: _ubicacion SQL devuelve solo ambiente (sin prefijo bd |)
REQ-SQL-01: gen_seccion_sql subtítulo incluye nombre de base_datos
REQ-CAT-01: _derive_component_config SQL → tipo_display usa inst.tipo o "SQL"
REQ-CAT-03: _derive_component_config liferay_build → tipo_display = "BUILD"
REQ-IMG-03/04: gen_seccion_api_iis/docker añade imagen si imagen_path está presente;
               no añade si es None; no lanza si el archivo no existe
"""

from __future__ import annotations

import logging
from unittest.mock import MagicMock, call, patch

import pytest
from docx import Document

from mgg_packify_api.routes.packages import _derive_component_config
from mgg_packify_api.schemas.package import ComponentIn, InstanceIn
from mgg_packify_api.services.doc_generator import (
    _ubicacion,
    gen_seccion_api_docker,
    gen_seccion_api_iis,
    gen_seccion_sql,
)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────


def _make_comp(tipo: str) -> ComponentIn:
    return ComponentIn(tipo=tipo)


def _make_inst(**kwargs) -> InstanceIn:
    return InstanceIn(**kwargs)


def _paragraphs_text(doc: Document) -> list[str]:
    """Devuelve el texto plano de todos los párrafos del doc."""
    return [p.text for p in doc.paragraphs]


# ─────────────────────────────────────────────────────────────────────────────
# Task 2.1 — _ubicacion SQL returns only ambiente (REQ-CAT-02)
# ─────────────────────────────────────────────────────────────────────────────


class TestUbicacionSql:
    """_ubicacion para tipo_clave='sql' debe retornar solo ambiente."""

    def _make_sql_comp(self, base_datos: str = "my_db") -> dict:
        return {"tipo_clave": "sql", "base_datos": base_datos}

    def _make_data(self, ambiente: str = "PRD", srv_bd: str = "") -> dict:
        return {"ambiente": ambiente, "servidores": {"bd": srv_bd}}

    def test_sql_ubicacion_returns_only_ambiente(self) -> None:
        """_ubicacion SQL sin servidor devuelve solo ambiente (REQ-CAT-02)."""
        comp = self._make_sql_comp("my_db")
        data = self._make_data("PRD", "")

        result = _ubicacion(comp, "PRD", data)

        assert result == "PRD"
        assert "my_db" not in result
        assert "|" not in result

    def test_sql_ubicacion_with_servidor_returns_only_ambiente_and_srv(self) -> None:
        """_ubicacion SQL con servidor devuelve ambiente (sin bd|) (REQ-CAT-02)."""
        comp = self._make_sql_comp("RAWRAPSIIF")
        data = self._make_data("QA", "sqlserver.internal")

        result = _ubicacion(comp, "QA", data)

        # Debe contener ambiente pero NO el nombre de bd ni el separador |
        assert "RAWRAPSIIF" not in result
        assert "|" not in result
        assert "QA" in result

    def test_sql_ubicacion_does_not_contain_bd_prefix(self) -> None:
        """El resultado de _ubicacion sql nunca empieza con nombre_bd (REQ-CAT-02)."""
        comp = self._make_sql_comp("SuperDB")
        data = self._make_data("PROD", "")

        result = _ubicacion(comp, "PROD", data)

        assert not result.startswith("SuperDB")


# ─────────────────────────────────────────────────────────────────────────────
# Task 2.3 — gen_seccion_sql h3 includes BD name (REQ-SQL-01)
# ─────────────────────────────────────────────────────────────────────────────


class TestGenSeccionSqlSubtitle:
    """gen_seccion_sql subtítulo debe incluir nombre de base_datos."""

    def test_subtitle_contains_bd_name(self) -> None:
        """h3 debe contener 'En la base de datos "mybd"' (REQ-SQL-01)."""
        doc = Document()
        gen_seccion_sql(
            doc,
            seccion_num="1",
            scripts=["init.sql"],
            base_datos="mybd",
            ambiente="PRD",
            servidor_bd="sqlserver.internal",
        )

        texts = _paragraphs_text(doc)
        combined = " ".join(texts)

        assert 'En la base de datos "mybd"' in combined

    def test_subtitle_contains_ambiente(self) -> None:
        """h3 también debe mencionar el ambiente (REQ-SQL-01)."""
        doc = Document()
        gen_seccion_sql(
            doc,
            seccion_num="2",
            scripts=["fix.sql"],
            base_datos="MYDB",
            ambiente="QA",
            servidor_bd="",
        )

        texts = _paragraphs_text(doc)
        combined = " ".join(texts)

        assert "QA" in combined

    def test_subtitle_format_full(self) -> None:
        """h3 sigue el formato exacto del diseño (REQ-SQL-01)."""
        doc = Document()
        gen_seccion_sql(
            doc,
            seccion_num="2",
            scripts=["a.sql"],
            base_datos="RAWRAPSIIF",
            ambiente="PRD",
            servidor_bd="",
        )

        texts = _paragraphs_text(doc)
        combined = " ".join(texts)

        expected_fragment = '2.1.- En la base de datos "RAWRAPSIIF" en el servidor de PRD:'
        assert expected_fragment in combined


# ─────────────────────────────────────────────────────────────────────────────
# Task 2.5 — _derive_component_config tipo_display (REQ-CAT-01, REQ-CAT-03)
# ─────────────────────────────────────────────────────────────────────────────


class TestDeriveComponentConfigTipoDisplay:
    """tipo_display para sql y liferay_build."""

    # REQ-CAT-01: SQL uses inst.tipo or "SQL" fallback
    def test_sql_tipo_display_uses_inst_tipo(self) -> None:
        """SQL con inst.tipo='DML' → tipo_display='DML' (REQ-CAT-01)."""
        comp = _make_comp("sql")
        inst = _make_inst(base_datos="DB", scripts=["a.sql"], tipo="DML")

        results = _derive_component_config(comp, inst)

        assert len(results) == 1
        assert results[0].tipo_display == "DML"

    def test_sql_tipo_display_fallback_to_SQL_when_empty(self) -> None:
        """SQL con inst.tipo='' → tipo_display='SQL' (REQ-CAT-01)."""
        comp = _make_comp("sql")
        inst = _make_inst(base_datos="DB", scripts=["a.sql"], tipo="")

        results = _derive_component_config(comp, inst)

        assert len(results) == 1
        assert results[0].tipo_display == "SQL"

    def test_sql_tipo_display_fallback_to_SQL_when_none(self) -> None:
        """SQL con inst.tipo=None → tipo_display='SQL' (REQ-CAT-01)."""
        comp = _make_comp("sql")
        inst = _make_inst(base_datos="DB", scripts=["a.sql"])
        # tipo defaults to "" in InstanceIn, but let's verify either way
        inst.tipo = None  # force None to cover the None branch

        results = _derive_component_config(comp, inst)

        assert len(results) == 1
        assert results[0].tipo_display == "SQL"

    # REQ-CAT-03: liferay_build tipo_display = "BUILD"
    def test_liferay_build_tipo_display_is_BUILD(self) -> None:
        """liferay_build → tipo_display='BUILD' (REQ-CAT-03)."""
        comp = _make_comp("liferay_build")
        inst = _make_inst(build_id="7957")

        results = _derive_component_config(comp, inst)

        assert len(results) == 1
        assert results[0].tipo_display == "BUILD"

    def test_liferay_build_tipo_display_is_not_Liferay(self) -> None:
        """liferay_build tipo_display NO debe ser 'Liferay' (REQ-CAT-03)."""
        comp = _make_comp("liferay_build")
        inst = _make_inst(build_id="1234")

        results = _derive_component_config(comp, inst)

        assert results[0].tipo_display != "Liferay"


# ─────────────────────────────────────────────────────────────────────────────
# Task 2.7 — gen_seccion_api_iis/docker add_picture behavior (REQ-IMG-03/04)
# ─────────────────────────────────────────────────────────────────────────────


class TestGenSeccionApiPicture:
    """gen_seccion_api_iis y gen_seccion_api_docker deben inyectar imagen si imagen_path está seteado."""

    # ── IIS ──────────────────────────────────────────────────────────────────

    def test_iis_with_imagen_path_calls_add_picture(self, tmp_path) -> None:
        """IIS con imagen_path válido → add_picture es llamado (REQ-IMG-03)."""
        img = tmp_path / "screenshot.png"
        img.write_bytes(b"fakepng")

        doc = Document()
        configs = [{"clave": "key", "valor": "val", "imagen_path": str(img)}]

        with patch.object(doc, "add_picture") as mock_pic:
            gen_seccion_api_iis(doc, "2", ["MySvc"], "QA", "srv", configs)

        mock_pic.assert_called_once()
        # Verificar que el primer argumento posicional es la ruta de la imagen
        assert mock_pic.call_args.args[0] == str(img)

    def test_iis_without_imagen_path_does_not_call_add_picture(self) -> None:
        """IIS con imagen_path=None → add_picture NO es llamado (REQ-IMG-03)."""
        doc = Document()
        configs = [{"clave": "key", "valor": "val", "imagen_path": None}]

        with patch.object(doc, "add_picture") as mock_pic:
            gen_seccion_api_iis(doc, "2", ["MySvc"], "QA", "srv", configs)

        mock_pic.assert_not_called()

    def test_iis_missing_file_logs_warning_and_no_raise(self, caplog) -> None:
        """IIS con imagen_path que no existe → warning loggeado, no exception (REQ-IMG-04)."""
        doc = Document()
        configs = [{"clave": "key", "valor": "val", "imagen_path": "/nonexistent/path/img.png"}]

        with caplog.at_level(logging.WARNING):
            # Must NOT raise
            gen_seccion_api_iis(doc, "2", ["MySvc"], "QA", "srv", configs)

        assert any("Imagen no encontrada" in r.message for r in caplog.records)

    def test_iis_config_without_imagen_path_key_does_not_call_add_picture(self) -> None:
        """IIS con config sin clave imagen_path → add_picture NO es llamado."""
        doc = Document()
        configs = [{"clave": "key", "valor": "val"}]

        with patch.object(doc, "add_picture") as mock_pic:
            gen_seccion_api_iis(doc, "2", ["MySvc"], "QA", "srv", configs)

        mock_pic.assert_not_called()

    # ── Docker ───────────────────────────────────────────────────────────────

    def test_docker_with_imagen_path_calls_add_picture(self, tmp_path) -> None:
        """Docker con imagen_path válido → add_picture es llamado (REQ-IMG-03)."""
        img = tmp_path / "diagram.png"
        img.write_bytes(b"fakepng")

        doc = Document()
        configs = [{"clave": "env", "valor": "prod", "imagen_path": str(img)}]

        with patch.object(doc, "add_picture") as mock_pic:
            gen_seccion_api_docker(
                doc, "3", [{"nombre": "DockerSvc", "jenkins": True, "actualizar_apim": True}],
                "PRD", "srv", configs,
            )

        mock_pic.assert_called_once()

    def test_docker_without_imagen_path_does_not_call_add_picture(self) -> None:
        """Docker con imagen_path=None → add_picture NO es llamado (REQ-IMG-03)."""
        doc = Document()
        configs = [{"clave": "env", "valor": "prod", "imagen_path": None}]

        with patch.object(doc, "add_picture") as mock_pic:
            gen_seccion_api_docker(
                doc, "3", [{"nombre": "DockerSvc", "jenkins": True, "actualizar_apim": True}],
                "PRD", "srv", configs,
            )

        mock_pic.assert_not_called()

    def test_docker_missing_file_logs_warning_and_no_raise(self, caplog) -> None:
        """Docker con imagen_path que no existe → warning, no exception (REQ-IMG-04)."""
        doc = Document()
        configs = [{"clave": "env", "valor": "prod", "imagen_path": "/does/not/exist.png"}]

        with caplog.at_level(logging.WARNING):
            gen_seccion_api_docker(
                doc, "3", [{"nombre": "DockerSvc", "jenkins": True, "actualizar_apim": True}],
                "PRD", "srv", configs,
            )

        assert any("Imagen no encontrada" in r.message for r in caplog.records)


# ─────────────────────────────────────────────────────────────────────────────
# REQ-DOC-01: Booleanos jenkins y actualizar_apim en gen_seccion_api_docker
# ─────────────────────────────────────────────────────────────────────────────


class TestGenSeccionApiDockerBooleanos:
    """Verifica que los booleanos jenkins y actualizar_apim controlan la presencia de pasos en el docx."""

    def test_jenkins_false_omits_jenkins_step(self) -> None:
        """jenkins=False → 'Hacer el despliegue CI/CD en Jenkins' NO aparece en el docx (REQ-DOC-01)."""
        doc = Document()
        apis = [{"nombre": "MiServicio", "jenkins": False, "actualizar_apim": True}]

        gen_seccion_api_docker(doc, "2", apis, "QA", "srv")

        texts = _paragraphs_text(doc)
        combined = " ".join(texts)
        assert "Jenkins" not in combined
        assert "Actualizar el schema del Api Management de AZURE" in combined

    def test_actualizar_apim_false_omits_apim_step(self) -> None:
        """actualizar_apim=False → 'Actualizar el schema del Api Management de AZURE' NO aparece (REQ-DOC-01)."""
        doc = Document()
        apis = [{"nombre": "MiServicio", "jenkins": True, "actualizar_apim": False}]

        gen_seccion_api_docker(doc, "2", apis, "QA", "srv")

        texts = _paragraphs_text(doc)
        combined = " ".join(texts)
        assert "Jenkins" in combined
        assert "Api Management de AZURE" not in combined

    def test_both_true_includes_both_steps(self) -> None:
        """jenkins=True y actualizar_apim=True → ambos pasos aparecen (retrocompatibilidad)."""
        doc = Document()
        apis = [{"nombre": "MiServicio", "jenkins": True, "actualizar_apim": True}]

        gen_seccion_api_docker(doc, "2", apis, "QA", "srv")

        texts = _paragraphs_text(doc)
        combined = " ".join(texts)
        assert "Jenkins" in combined
        assert "Api Management de AZURE" in combined

    def test_both_false_omits_both_steps(self) -> None:
        """jenkins=False y actualizar_apim=False → ningún paso opcional aparece."""
        doc = Document()
        apis = [{"nombre": "MiServicio", "jenkins": False, "actualizar_apim": False}]

        gen_seccion_api_docker(doc, "2", apis, "QA", "srv")

        texts = _paragraphs_text(doc)
        combined = " ".join(texts)
        assert "Jenkins" not in combined
        assert "Api Management de AZURE" not in combined
        # Pero el nombre del servicio sí debe aparecer
        assert "MiServicio" in combined

    def test_default_true_when_key_absent(self) -> None:
        """Cuando la key jenkins/actualizar_apim está ausente, default=True (retrocompatibilidad)."""
        doc = Document()
        # Dict sin las keys de booleanos — simula paquete generado antes del cambio
        apis = [{"nombre": "ServicioViejo"}]

        gen_seccion_api_docker(doc, "2", apis, "QA", "srv")

        texts = _paragraphs_text(doc)
        combined = " ".join(texts)
        assert "Jenkins" in combined
        assert "Api Management de AZURE" in combined

    def test_multiple_instances_independent_flags(self) -> None:
        """Múltiples instancias con flags distintos → cada una aplica sus propios flags."""
        doc = Document()
        apis = [
            {"nombre": "SvcA", "jenkins": True, "actualizar_apim": False},
            {"nombre": "SvcB", "jenkins": False, "actualizar_apim": True},
        ]

        gen_seccion_api_docker(doc, "2", apis, "QA", "srv")

        texts = _paragraphs_text(doc)
        combined = " ".join(texts)
        # Jenkins debe aparecer (por SvcA)
        assert "Jenkins" in combined
        # APIM debe aparecer (por SvcB)
        assert "Api Management de AZURE" in combined
        # Ambos servicios referenciados
        assert "SvcA" in combined
        assert "SvcB" in combined
