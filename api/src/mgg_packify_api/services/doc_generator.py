"""
Generador de documentos .docx para packages de instalación.

Port 1:1 de la lógica de generación de v1 (package_generator.py).
La función principal es `generate_document(config, output_path)`.

CRÍTICO — reglas de fidelidad con v1:
- add_footer() usa XML directo con tabs en posición 4419/8838 — copiado literalmente
- build_components_table() usa gridSpan + XML para merge de celdas — copiado literalmente
- QA → UAT en gen_seccion_liferay_build:
  `ambiente_display = "UAT" if ambiente == "QA" else ambiente`
- Autor footer hardcodeado: "Manuel García González"
- Colores: header 70AD47, border C5E0B3, group separator E2EFD9
"""

from __future__ import annotations

import logging
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from mgg_packify_api.services.component import ETIQUETA_GRUPO, ComponentType
from mgg_packify_api.services.package import PackageConfig

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTES DE ESTILO (extraídas del XML real de los manuales existentes)
# ─────────────────────────────────────────────────────────────────────────────

COLOR_VERDE_HEADER = "70AD47"   # Verde encabezados y fila header de tabla
COLOR_VERDE_BORDES = "C5E0B3"   # Verde claro para bordes de tabla
COLOR_VERDE_GRUPO  = "E2EFD9"   # Verde muy claro: fila separadora de grupo
COLOR_BLANCO       = "FFFFFF"
COLOR_NEGRO        = "000000"
FUENTE_DEFAULT     = "Calibri"

AUTOR_FOOTER = "Manuel García González"

# ─────────────────────────────────────────────────────────────────────────────
# FUNCIÓN PRINCIPAL
# ─────────────────────────────────────────────────────────────────────────────


def generate_document(config: PackageConfig, output_path: Path) -> None:
    """
    Genera el .docx del manual de instalación.

    Args:
        config: Configuración completa del package.
        output_path: Ruta completa donde guardar el archivo .docx.
    """
    # Construir el dict `data` en el formato exacto que usa v1
    servidores = {
        "api":     config.servidores.api,
        "bd":      config.servidores.bd,
        "blob":    config.servidores.blob,
        "liferay": config.servidores.liferay,
    }

    # Convertir ComponentConfig a dicts del formato v1
    componentes_v1 = []
    for comp in config.componentes:
        d: dict = {
            "tipo_clave":    comp.tipo_clave.value,
            "nombre_display": comp.nombre_display,
            "tipo_display":   comp.tipo_display,
            "contenedor":     comp.contenedor,
            "estatus":        comp.estatus,
            "scripts":        comp.scripts,
            "base_datos":     comp.base_datos,
            "nombre_servicio": comp.nombre_servicio,
            "configs":        comp.configs,
            "archivos":       comp.archivos,
            "build_id":       comp.build_id,
            "nombre":         comp.nombre,
            "tipo":           comp.tipo,
            "crear_pagina":   comp.crear_pagina,
            "pagina":         comp.pagina,
            "widgets":        comp.widgets,
        }
        componentes_v1.append(d)

    data = {
        "ticket":     config.ticket,
        "hu_nombre":  config.hu_nombre,
        "ambiente":   config.ambiente,
        "iteracion":  config.iteracion,
        "componentes": componentes_v1,
        "servidores":  servidores,
    }

    generar_manual(data, str(output_path))


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS DE FORMATO DOCX (port 1:1 de v1)
# ─────────────────────────────────────────────────────────────────────────────


def set_cell_background(cell, hex_color: str):
    """Aplica color de fondo a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_table_borders(table, color: str):
    """Aplica bordes de color a toda la tabla."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tblBorders.append(border)
    # Quitar tblBorders anterior si existe
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _override_heading_style(paragraph, size_pt: int, color_hex: str,
                             bold: bool = True, indent_cm: float = 0.0):
    """
    Fuerza color, fuente y tamaño en el rPr del estilo Heading para que
    siempre se vea igual independientemente del tema del documento.
    """
    for run in paragraph.runs:
        run.font.name = FUENTE_DEFAULT
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.font.color.rgb = RGBColor.from_string(color_hex)
    if indent_cm:
        paragraph.paragraph_format.left_indent = Cm(indent_cm)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)


def add_paragraph(doc, text: str, bold=False, size=11, color=None,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    """Agrega un párrafo con formato."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FUENTE_DEFAULT
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_heading_green(doc, text: str):
    """Banner verde centrado del título del manual (no es un Heading de jerarquía)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), COLOR_VERDE_HEADER)
    pPr.append(shd)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FUENTE_DEFAULT
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(COLOR_BLANCO)
    return p


def add_h1(doc, text: str):
    """
    Heading 1 — secciones principales: 'Componentes afectados', 'Proceso'.
    Verde #70AD47, bold, 12pt.
    """
    p = doc.add_heading(text, level=1)
    _override_heading_style(p, size_pt=12, color_hex=COLOR_VERDE_HEADER, bold=True, indent_cm=0)
    return p


def add_h2(doc, text: str):
    """
    Heading 2 — sub-secciones: '2.1.- API', '2.2.- SQL'.
    Negro, bold, 11pt, indent 0.5 cm.
    """
    p = doc.add_heading(text, level=2)
    _override_heading_style(p, size_pt=11, color_hex=COLOR_NEGRO, bold=True, indent_cm=0.5)
    return p


def add_h3(doc, text: str):
    """
    Heading 3 — primer paso por sub-sección.
    Negro, normal weight, 11pt, indent 1.0 cm.
    """
    p = doc.add_heading(text, level=3)
    _override_heading_style(p, size_pt=11, color_hex=COLOR_NEGRO, bold=False, indent_cm=1.0)
    return p


# Alias de compatibilidad
def add_section_heading(doc, text: str):
    """Alias → add_h1."""
    return add_h1(doc, text)


def add_child(doc, text: str, bold_keywords=None, indent_cm: float = 1.5):
    """
    Párrafo hijo bajo un Heading 3.
    Sin número de paso — texto plano con indentación, negro, 11pt.
    Acepta bold_keywords para resaltar fragmentos específicos.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(indent_cm)
    if bold_keywords:
        _add_mixed_run(p, text, bold_keywords)
    else:
        run = p.add_run(text)
        run.font.name = FUENTE_DEFAULT
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(COLOR_NEGRO)
    return p


def _add_mixed_run(paragraph, text: str, bold_fragments: list):
    """Agrega texto con ciertos fragmentos en bold."""
    remaining = text
    for fragment in bold_fragments:
        idx = remaining.find(fragment)
        if idx == -1:
            continue
        if idx > 0:
            run = paragraph.add_run(remaining[:idx])
            run.font.name = FUENTE_DEFAULT
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor.from_string(COLOR_NEGRO)
        run_bold = paragraph.add_run(fragment)
        run_bold.bold = True
        run_bold.font.name = FUENTE_DEFAULT
        run_bold.font.size = Pt(11)
        run_bold.font.color.rgb = RGBColor.from_string(COLOR_NEGRO)
        remaining = remaining[idx + len(fragment):]
    if remaining:
        run = paragraph.add_run(remaining)
        run.font.name = FUENTE_DEFAULT
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(COLOR_NEGRO)


def add_note(doc, text: str):
    """Agrega una nota en itálica."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"NOTA: {text}")
    run.italic = True
    run.font.name = FUENTE_DEFAULT
    run.font.size = Pt(10)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# TABLA DE COMPONENTES (port 1:1 de v1)
# ─────────────────────────────────────────────────────────────────────────────


def _merge_row_horizontal(row, n_cols: int):
    """
    Combina todas las celdas de una fila en una sola (merge horizontal).
    Usa gridSpan en la primera celda y marca las demás con w:tcW w=0.
    """
    first_tc = row.cells[0]._tc
    tcPr = first_tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        first_tc.insert(0, tcPr)

    # Quitar gridSpan anterior si existe
    existing = tcPr.find(qn("w:gridSpan"))
    if existing is not None:
        tcPr.remove(existing)

    gs = OxmlElement("w:gridSpan")
    gs.set(qn("w:val"), str(n_cols))
    tcPr.append(gs)

    # Eliminar las celdas extra del XML (dejar solo la primera)
    tr = row._tr
    tcs = tr.findall(qn("w:tc"))
    for tc in tcs[1:]:
        tr.remove(tc)


def _add_group_separator_row(table, label: str, n_cols: int):
    """
    Agrega una fila separadora de grupo con:
    - celda única que abarca n_cols columnas
    - fondo verde claro E2EFD9
    - texto centrado, bold, 10pt, negro
    """
    row = table.add_row()
    cell = row.cells[0]

    # Fondo verde claro
    set_cell_background(cell, COLOR_VERDE_GRUPO)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(label)
    run.bold = True
    run.font.name = FUENTE_DEFAULT
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(COLOR_NEGRO)

    # Merge horizontal
    _merge_row_horizontal(row, n_cols)


def build_components_table(doc, componentes: list, ambiente: str):
    """
    Crea la tabla de componentes afectados con filas separadoras de grupo.

    componentes: lista de dicts con keys:
        nombre, estatus, tipo, contenedor, ubicacion, tipo_clave
    Los componentes del mismo grupo se agrupan bajo una fila separadora
    con el nombre del grupo (API, SQL, LIFERAY, etc.) en verde claro.
    """
    N_COLS = 5
    table = doc.add_table(rows=1, cols=N_COLS)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ── Fila header ──────────────────────────────────────────────────────────
    hdr_cells = table.rows[0].cells
    headers = ["NOMBRE", "ESTATUS", "TIPO", "CONTENEDOR", "UBICACIÓN"]
    for cell, header in zip(hdr_cells, headers):
        set_cell_background(cell, COLOR_VERDE_HEADER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.name = FUENTE_DEFAULT
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(COLOR_BLANCO)

    # ── Agrupar componentes por tipo_clave (respetando orden canónico) ────────
    # Normalizar: agrupar api_iis y api_docker bajo etiqueta "API" manteniendo orden
    grupos_vistos: list = []
    grupos_dict: dict = {}
    for comp in componentes:
        tk = comp.get("tipo_clave", "")
        # Convertir a ComponentType para usar ETIQUETA_GRUPO
        try:
            ct = ComponentType(tk)
            etiqueta = ETIQUETA_GRUPO.get(ct, tk.upper())
        except ValueError:
            etiqueta = tk.upper()
        # clave de agrupación visual (api_iis y api_docker comparten "API")
        clave_visual = etiqueta
        if clave_visual not in grupos_dict:
            grupos_dict[clave_visual] = []
            grupos_vistos.append(clave_visual)
        grupos_dict[clave_visual].append(comp)

    # ── Escribir grupos en la tabla ──────────────────────────────────────────
    for etiqueta in grupos_vistos:
        # Fila separadora de grupo
        _add_group_separator_row(table, etiqueta, N_COLS)

        # Filas de datos del grupo
        for comp in grupos_dict[etiqueta]:
            row_cells = table.add_row().cells
            values = [
                comp.get("nombre", ""),
                comp.get("estatus", "modificado"),
                comp.get("tipo", ""),
                comp.get("contenedor", ""),
                comp.get("ubicacion", ambiente),
            ]
            for i, (cell, value) in enumerate(zip(row_cells, values)):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                # NOMBRE (col 0) queda LEFT, el resto centrado
                if i > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(value)
                run.font.name = FUENTE_DEFAULT
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor.from_string(COLOR_NEGRO)

    set_table_borders(table, COLOR_VERDE_BORDES)
    return table


# ─────────────────────────────────────────────────────────────────────────────
# GENERADORES DE SECCIONES POR TIPO (port 1:1 de v1)
# ─────────────────────────────────────────────────────────────────────────────


def gen_seccion_sql(doc, seccion_num: str, scripts: list, base_datos: str,
                    ambiente: str, servidor_bd: str):
    """Genera la sección de proceso para scripts SQL."""
    add_h2(doc, f"{seccion_num}.- SQL")

    # REQ-SQL-01: incluir nombre de base_datos en el subtítulo
    add_h3(doc, f'{seccion_num}.1.- En la base de datos "{base_datos}" en el servidor de {ambiente}:')

    for script in scripts:
        add_child(doc,
                  f'Ejecutar el script "{script}"',
                  bold_keywords=[f'"{script}"'])


def gen_seccion_api_iis(doc, seccion_num: str, apis: list, ambiente: str,
                         servidor_api: str, configs: list | None = None):
    """Genera la sección de proceso para APIs en IIS (.zip)."""
    add_h2(doc, f"{seccion_num}.- API")

    add_h3(doc, f"{seccion_num}.1.-  En el servidor de servicios en {ambiente}:")

    for api in apis:
        zip_name = api if api.endswith(".zip") else f"{api}.zip"
        add_child(doc,
                  f'Actualizar el servicio "{api}" con el contenido del zip "{zip_name}"',
                  bold_keywords=[f'"{api}"', f'"{zip_name}"'])

    if configs:
        add_child(
            doc,
            "Actualizar las siguientes configuraciones en el archivo"
            " de configuración del servicio:",
        )
        for cfg in configs:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(2.5)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"• {cfg['clave']}: {cfg['valor']}")
            run.font.name = FUENTE_DEFAULT
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor.from_string(COLOR_NEGRO)
            # REQ-IMG-03/04: inyectar imagen si imagen_path está presente
            if cfg.get('imagen_path'):
                try:
                    doc.add_picture(cfg['imagen_path'], width=Cm(14))
                except Exception as e:
                    logging.warning(f"Imagen no encontrada: {cfg['imagen_path']} — {e}")


def gen_seccion_api_docker(doc, seccion_num: str, apis: list, ambiente: str,
                            servidor_api: str, configs: list | None = None):
    """Genera la sección de proceso para APIs con pipeline Docker / Jenkins CI-CD."""
    add_h2(doc, f"{seccion_num}.- API")

    add_h3(doc, f"{seccion_num}.1.-  En el servidor de servicios en {ambiente}:")

    for api in apis:
        add_child(doc, f'Actualizar el servicio "{api}"', bold_keywords=[f'"{api}"'])
        add_child(doc, "Hacer el despliegue CI/CD en Jenkins", bold_keywords=["Jenkins"])
        add_child(doc, "Actualizar el schema del Api Management de AZURE",
                  bold_keywords=["Api Management de AZURE"])

    if configs:
        add_child(doc, "Actualizar las siguientes variables de entorno/configuración:")
        for cfg in configs:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(2.5)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"• {cfg['clave']}: {cfg['valor']}")
            run.font.name = FUENTE_DEFAULT
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor.from_string(COLOR_NEGRO)
            # REQ-IMG-03/04: inyectar imagen si imagen_path está presente
            if cfg.get('imagen_path'):
                try:
                    doc.add_picture(cfg['imagen_path'], width=Cm(14))
                except Exception as e:
                    logging.warning(f"Imagen no encontrada: {cfg['imagen_path']} — {e}")


def gen_seccion_blob_storage(doc, seccion_num: str, blobs: list, ambiente: str,
                              servidor_blob: str):
    """Genera la sección de proceso para Azure Blob Storage."""
    add_h2(doc, f"{seccion_num}.- Blob Storage")

    add_h3(doc, f"{seccion_num}.1.-  En el servidor de Azure Blob Storage "
                f"de {ambiente}:")

    # Agrupar por carpeta de blob
    carpetas: dict = {}
    for blob in blobs:
        carpeta = blob.get("carpeta", blob["nombre"].rsplit(".", 1)[0])
        if carpeta not in carpetas:
            carpetas[carpeta] = []
        carpetas[carpeta].append(blob)

    for carpeta, archivos in carpetas.items():
        add_child(doc,
                  f'Validar si existe la carpeta "{carpeta}" (si no existe, crearla)',
                  bold_keywords=[f'"{carpeta}"'])
        add_child(doc,
                  f'Dentro de la carpeta "{carpeta}" subir los siguientes archivos:',
                  bold_keywords=[f'"{carpeta}"'])

        for archivo in archivos:
            nombre = archivo["nombre"]
            # bullet del archivo — indent extra
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(2.5)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"• {nombre}")
            run.font.name = FUENTE_DEFAULT
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor.from_string(COLOR_NEGRO)

            # Sub-línea URL SAS
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(3.5)
            p2.paragraph_format.space_after = Pt(2)
            run2 = p2.add_run("Generar URL SAS")
            run2.font.name = FUENTE_DEFAULT
            run2.font.size = Pt(11)
            run2.italic = True
            run2.font.color.rgb = RGBColor.from_string(COLOR_NEGRO)


def gen_seccion_liferay_build(doc, seccion_num: str, build_id: str, ambiente: str):
    """Genera la sección de proceso para deploy de build de Liferay.

    CRÍTICO: QA → UAT en ambiente_display.
    """
    add_h2(doc, f"{seccion_num}.- Liferay")

    ambiente_display = "UAT" if ambiente == "QA" else ambiente
    add_h3(doc, f"{seccion_num}.1.-  Hacer deploy de Liferay en ambiente "
                f"{ambiente_display} la build # {build_id}")


def gen_seccion_liferay(doc, seccion_num: str, remote_apps: list, ambiente: str,
                         servidor_liferay: str):
    """Genera la sección de proceso para Liferay."""
    add_h2(doc, f"{seccion_num}.- Liferay")

    add_h3(doc, f"{seccion_num}.1.-  En el servidor de Liferay de "
                f"{ambiente}:")

    for app in remote_apps:
        nombre_app = app["nombre"]
        js_url_var = f"URL del Blob Storage del archivo {nombre_app}.js"
        css_url_var = f"URL del Blob Storage del archivo {nombre_app}.css"

        if app.get("tipo") == "remote_app":
            add_child(
                doc,
                f'Crear (o actualizar) la Remote App "{nombre_app}"'
                " con los siguientes campos:",
                bold_keywords=[f'"{nombre_app}"'])
            for linea in [f"URL: {js_url_var}", f"CSS: {css_url_var}"]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(2.5)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(f"• {linea}")
                run.font.name = FUENTE_DEFAULT
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor.from_string(COLOR_NEGRO)

        if app.get("crear_pagina"):
            pagina = app.get("pagina", nombre_app)
            add_child(doc, "Dirigirse al Sitio RETAIL → Páginas Privadas",
                      bold_keywords=["RETAIL", "Páginas Privadas"])
            add_child(doc, f'Crear la página "{pagina}"', bold_keywords=[f'"{pagina}"'])
            add_child(doc, "Editar la página → pestaña Widgets / Remote Apps")
            for widget in app.get("widgets", [nombre_app]):
                add_child(doc, f'Arrastrar la Remote App "{widget}"',
                          bold_keywords=[f'"{widget}"'])
            add_child(doc, "Dar clic en Publish", bold_keywords=["Publish"])


def gen_seccion_assets(doc, seccion_num: str, assets: list, ambiente: str,
                        servidor_liferay: str):
    """Genera la sección para subir assets a Liferay Documents and Media."""
    add_h2(doc, f"{seccion_num}.- Assets (Liferay Documents and Media)")

    add_h3(doc, f"{seccion_num}.1.-  En el servidor de Liferay de "
                f"{ambiente}:")

    add_child(doc, "Dirigirse a Content & Data → Documents and Media",
              bold_keywords=["Content & Data", "Documents and Media"])

    for asset in assets:
        add_child(doc, f'Subir el archivo "{asset}"', bold_keywords=[f'"{asset}"'])


def gen_seccion_azure_apim(doc, seccion_num: str, servicios: list, ambiente: str):
    """Genera la sección para Azure API Management."""
    add_h2(doc, f"{seccion_num}.- Azure API Management")

    add_h3(doc, f"{seccion_num}.1.-  En Azure API Management de {ambiente}:")

    for svc in servicios:
        add_child(doc, f'Actualizar el servicio "{svc}"', bold_keywords=[f'"{svc}"'])


# ─────────────────────────────────────────────────────────────────────────────
# FOOTER (port 1:1 de v1 — XML directo con tabs en posición 4419/8838)
# ─────────────────────────────────────────────────────────────────────────────

MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}


def fecha_larga(d: date) -> str:
    """Devuelve la fecha en formato '31 de marzo del 2026'."""
    return f"{d.day} de {MESES_ES[d.month]} del {d.year}"


def add_footer(doc, autor: str):
    """
    Agrega pie de página con:
      izquierda → 'Autor: {autor}'
      derecha   → fecha actual en formato '31 de marzo del 2026'
    Replica el estilo del footer de los manuales originales:
    tabs en posición central (4419) y derecha (8838), color negro, Calibri 9pt.

    CRÍTICO: Usa XML directo — NO modificar sin revisar el XML de v1.
    """
    fecha = fecha_larga(date.today())

    section = doc.sections[0]
    section.footer_distance = Cm(1.25)
    footer = section.footer

    # Limpiar párrafos por defecto
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    # Crear párrafo del footer con XML directo para replicar el estilo exacto
    ftr = footer._element
    p_elem = OxmlElement("w:p")

    # pPr: sin bordes, tabs centro/derecha, sin espacio después
    pPr = OxmlElement("w:pPr")

    pBdr = OxmlElement("w:pBdr")
    for lado in ["top", "left", "bottom", "right", "between"]:
        b = OxmlElement(f"w:{lado}")
        b.set(qn("w:val"), "nil")
        pBdr.append(b)
    pPr.append(pBdr)

    tabs = OxmlElement("w:tabs")
    tab_center = OxmlElement("w:tab")
    tab_center.set(qn("w:val"), "center")
    tab_center.set(qn("w:pos"), "4419")
    tabs.append(tab_center)
    tab_right = OxmlElement("w:tab")
    tab_right.set(qn("w:val"), "right")
    tab_right.set(qn("w:pos"), "8838")
    tabs.append(tab_right)
    pPr.append(tabs)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

    rPr_p = OxmlElement("w:rPr")
    color_p = OxmlElement("w:color")
    color_p.set(qn("w:val"), "000000")
    rPr_p.append(color_p)
    pPr.append(rPr_p)

    p_elem.append(pPr)

    def make_run(text, is_tab=False):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        rPr.append(color)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "18")   # 9pt = 18 half-points
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "18")
        rPr.append(szCs)
        r.append(rPr)
        if is_tab:
            r.append(OxmlElement("w:tab"))
        else:
            t = OxmlElement("w:t")
            t.text = text
            if text.startswith(" ") or text.endswith(" "):
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
        return r

    # Izquierda: "Autor: Manuel García González"
    p_elem.append(make_run(f"Autor: {autor}"))
    # Tab al centro
    p_elem.append(make_run("", is_tab=True))
    # Tab a la derecha
    p_elem.append(make_run("", is_tab=True))
    # Derecha: fecha
    p_elem.append(make_run(fecha))

    ftr.append(p_elem)


# ─────────────────────────────────────────────────────────────────────────────
# HELPER DE UBICACIÓN (port 1:1 de v1)
# ─────────────────────────────────────────────────────────────────────────────


def _ubicacion(comp: dict, ambiente: str, data: dict) -> str:
    """Determina el texto de ubicación para la tabla de componentes."""
    tipo = comp["tipo_clave"]
    if tipo == "sql":
        # REQ-CAT-02: solo mostrar ambiente, sin prefijo de base_datos
        return ambiente
    elif tipo in ("api_iis", "api_docker", "apim"):
        srv = data["servidores"].get("api", "")
        return f"{ambiente}: {srv}" if srv else ambiente
    elif tipo == "blob":
        srv = data["servidores"].get("blob", "")
        return f"{ambiente}: {srv}" if srv else ambiente
    elif tipo in ("liferay", "liferay_build", "assets"):
        srv = data["servidores"].get("liferay", "")
        return f"{ambiente}: {srv}" if srv else ambiente
    return ambiente


# ─────────────────────────────────────────────────────────────────────────────
# FUNCIÓN PRINCIPAL INTERNA (port 1:1 de v1)
# ─────────────────────────────────────────────────────────────────────────────


def generar_manual(data: dict, output_path: str):
    """
    Genera el .docx del manual de instalación.

    data: dict con toda la información recopilada (formato v1).
    output_path: ruta completa donde guardar el archivo.

    Esta función es un port 1:1 de v1. NO modificar el comportamiento.
    """
    doc = Document()

    # ── Márgenes ─────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    ambiente  = data["ambiente"]
    ticket    = data["ticket"]
    hu_nombre = data["hu_nombre"]

    # ── Título ────────────────────────────────────────────────────────────────
    add_paragraph(doc, "Manual de instalación", bold=True, size=16, space_after=4,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading_green(doc, f"{ticket} - Portal Retail - {hu_nombre}")
    doc.add_paragraph()  # espacio

    # ── Sección 1: Componentes afectados (H1) ────────────────────────────────
    add_h1(doc, "1.- Componentes afectados:")

    # Construir filas de la tabla desde los componentes
    tabla_rows = []
    for comp in data["componentes"]:
        tabla_rows.append({
            "nombre":     comp["nombre_display"],
            "estatus":    comp["estatus"],
            "tipo":       comp["tipo_display"],
            "contenedor": comp["contenedor"],
            "ubicacion":  _ubicacion(comp, ambiente, data),
            "tipo_clave": comp["tipo_clave"],
        })

    build_components_table(doc, tabla_rows, ambiente)
    doc.add_paragraph()

    # Nota estándar
    add_note(doc,
        "Los componentes relacionados en este listado deberán ser respaldados y actualizados "
        "en caso de que existan, y en caso de no existir deberán ser agregados "
        "en la ruta especificada.")
    doc.add_paragraph()

    # ── Sección 2: Proceso (H1) ──────────────────────────────────────────────
    add_h1(doc, "2.- Proceso")

    # seccion arranca en 1 → genera "2.1.- X", "2.2.- Y", etc.
    seccion = 1
    seccion_padre = "2"   # número del H1 "Proceso"
    # Orden canónico de sub-secciones dentro de Proceso
    orden = ["liferay_build", "sql", "api_iis", "api_docker", "blob", "liferay", "assets", "apim"]

    for tipo_clave in orden:
        grupo = [c for c in data["componentes"] if c["tipo_clave"] == tipo_clave]
        if not grupo:
            continue

        # prefijo compuesto: "2.1", "2.2", etc.
        prefix = f"{seccion_padre}.{seccion}"

        if tipo_clave == "liferay_build":
            for c in grupo:
                gen_seccion_liferay_build(doc, prefix, c["build_id"], ambiente)
                seccion += 1
                prefix = f"{seccion_padre}.{seccion}"

        elif tipo_clave == "sql":
            # Agrupar por base de datos
            por_bd: dict = {}
            for c in grupo:
                bd = c["base_datos"]
                if bd not in por_bd:
                    por_bd[bd] = []
                por_bd[bd].extend(c["scripts"])
            for bd, scripts in por_bd.items():
                gen_seccion_sql(doc, prefix, scripts, bd, ambiente,
                                data["servidores"].get("bd", f"Servidor BD {ambiente}"))
                seccion += 1
                prefix = f"{seccion_padre}.{seccion}"

        elif tipo_clave == "api_iis":
            nombres = [c["nombre_servicio"] for c in grupo]
            configs: list = []
            for c in grupo:
                configs.extend(c.get("configs", []))
            gen_seccion_api_iis(doc, prefix, nombres, ambiente,
                                data["servidores"].get("api", f"Servidor de servicios {ambiente}"),
                                configs if configs else None)
            seccion += 1

        elif tipo_clave == "api_docker":
            nombres = [c["nombre_servicio"] for c in grupo]
            configs = []
            for c in grupo:
                configs.extend(c.get("configs", []))
            servidor_api = data["servidores"].get(
                "api", f"Servidor de servicios {ambiente}"
            )
            gen_seccion_api_docker(doc, prefix, nombres, ambiente,
                                   servidor_api,
                                   configs if configs else None)
            seccion += 1

        elif tipo_clave == "blob":
            archivos: list = []
            for c in grupo:
                archivos.extend(c["archivos"])
            gen_seccion_blob_storage(doc, prefix, archivos, ambiente,
                                     data["servidores"].get("blob", f"Servidor Azure {ambiente}"))
            seccion += 1

        elif tipo_clave == "liferay":
            gen_seccion_liferay(doc, prefix, grupo, ambiente,
                                data["servidores"].get("liferay", f"Servidor Liferay {ambiente}"))
            seccion += 1

        elif tipo_clave == "assets":
            archivos = []
            for c in grupo:
                archivos.extend(c["archivos"])
            gen_seccion_assets(doc, prefix, archivos, ambiente,
                               data["servidores"].get("liferay", f"Servidor Liferay {ambiente}"))
            seccion += 1

        elif tipo_clave == "apim":
            servicios = [c["nombre_servicio"] for c in grupo]
            gen_seccion_azure_apim(doc, prefix, servicios, ambiente)
            seccion += 1

    add_footer(doc, AUTOR_FOOTER)
    doc.save(output_path)
